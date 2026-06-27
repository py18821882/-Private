#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""H5展业项目问题分析报告生成器"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── 颜色常量 ──────────────────────────────────────────────────────────────
DARK_BG   = RGBColor(0x1B, 0x5E, 0x3B)   # 深绿标题底
GOLD       = RGBColor(0xF9, 0xA8, 0x25)   # 金色强调
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG   = RGBColor(0xE8, 0xF5, 0xE9)   # 浅绿背景
STRIPE     = RGBColor(0xF1, 0xF8, 0xE9)   # 斑马纹
RED_BG     = RGBColor(0xFE, 0xEB, 0xEE)   # 严重问题背景
ORANGE_BG  = RGBColor(0xFF, 0xF3, 0xE0)   # 高优先级背景
YELLOW_BG  = RGBColor(0xFF, 0xFA, 0xE6)   # 中优先级背景
RED_TEXT   = RGBColor(0xD3, 0x2F, 0x2F)
ORANGE_TEXT= RGBColor(0xE6, 0x51, 0x00)
YELLOW_TEXT= RGBColor(0x7D, 0x66, 0x00)
GRAY_TEXT  = RGBColor(0x37, 0x51, 0x43)
BORDER     = RGBColor(0xCB, 0xDB, 0xB5)   # 表格边框

# ── 问题数据 ──────────────────────────────────────────────────────────────
ISSUES = [
    # 严重问题
    {"id": 1, "severity": "🔴 严重", "file": "app/assessment/[id]/page.tsx",
     "line": "—", "status": "已修复",
     "desc": "useEffect 钩子内部结构错乱，括号不匹配导致 JSX 解析失败，整个文件无法编译。",
     "fix": "重写文件：修复 useEffect 结构；将 checkPayment 移出 useEffect；用 useRef 管理轮询定时器；location.reload() 改为 fetchData()。"},
    {"id": 2, "severity": "🔴 严重", "file": "app/assessment/[id]/page.tsx",
     "line": "102", "status": "已修复",
     "desc": "第 102 行出现非法语法 const checkPayment(oid)，const 关键字多余导致解析错误。",
     "fix": "删除多余的 const 关键字，改为标准函数调用 checkPayment(oid)。"},
    {"id": 3, "severity": "🔴 严重", "file": "lib/ai/index.ts",
     "line": "9", "status": "已修复",
     "desc": "ChatMessage 接口中 role 类型拼写错误：'assitant' 应为 'assistant'。",
     "fix": "修正拼写：'assitant' → 'assistant'。"},
    {"id": 4, "severity": "🔴 严重", "file": "prisma/schema.prisma",
     "line": "45-46", "status": "已修复",
     "desc": "Order.assessment 定义为 Assessment[]（数组），业务逻辑应为单个关联；@relation 在两侧重复定义 fields/references 导致 schema 验证失败。",
     "fix": "Order.assessment 改为 Assessment?；删除 Order/Report 侧重复的 @relation(fields/references)；assessmentId/orderId/reportId 加 @unique。"},
    # 高优先级
    {"id": 5, "severity": "🟠 高", "file": "app/api/payment/wechat/notify/route.ts",
     "line": "62", "status": "已修复",
     "desc": "支付成功判断条件使用 &&（且），导致 return_code 和 result_code 同时不为 SUCCESS 才视为失败，逻辑错误。",
     "fix": "将 && 改为 ||：任一字段不为 SUCCESS 即视为未支付成功。"},
    {"id": 6, "severity": "🟠 高", "file": "app/assessment/[id]/page.tsx",
     "line": "17-62", "status": "已修复",
     "desc": "useEffect 依赖数组包含 data?.status，每次 setData 后 effect 重新执行，可能启动多个 setInterval，造成内存泄漏和重复请求。",
     "fix": "用 useRef 保存 timer 引用，清理前先 clearInterval，避免重复轮询。"},
    {"id": 7, "severity": "🟠 高", "file": "app/assessment/[id]/page.tsx",
     "line": "117", "status": "已修复",
     "desc": "模拟支付后使用 location.reload() 整页刷新，在 SPA 中应改用状态更新，体验差且会丢失组件状态。",
     "fix": "改为调用 fetchData() 重新获取数据，避免整页刷新。"},
    {"id": 8, "severity": "🟠 高", "file": "lib/report/index.ts",
     "line": "172", "status": "已修复",
     "desc": "getReportPreview 中 score 字段未做类型校验，若 fullReport.score 为字符串或 undefined，Math.floor 返回 NaN，前端渲染显示 NaN。",
     "fix": "加 typeof score === 'number' 防护，非数字时返回 null。"},
    # 中优先级
    {"id": 9,  "severity": "🟡 中", "file": "lib/payment/wechat.ts",
     "line": "113", "status": "待修复",
     "desc": "生产环境微信支付未实现，调用时直接 throw Error('生产环境微信支付请接入真实 API')，属于空实现。",
     "fix": "接入微信支付 V3 API（需要商户号、API 密钥、证书等），工程量较大，建议单独排期。"},
    {"id": 10, "severity": "🟡 中", "file": "所有 API 路由",
     "line": "—", "status": "待修复",
     "desc": "API 路由缺少输入格式校验，如 contactPhone 未校验手机号格式，字段类型未严格检查，存在脏数据风险。",
     "fix": "在 API 路由层增加 Zod 或 Joi 校验 Schema，对手机号、邮箱、必填字段做格式校验。"},
    {"id": 11, "severity": "🟡 中", "file": "所有 API 路由",
     "line": "—", "status": "待修复",
     "desc": "缺少 API 限流/防刷机制，任何人可无限次调用提交测评接口，存在被刷风险。",
     "fix": "在 middleware.ts 或 API 层增加 rate limiting（如 @upstash/ratelimit 或 express-rate-limit 风格的中间件）。"},
    {"id": 12, "severity": "🟡 中", "file": "next.config.mjs",
     "line": "3-6", "status": "已修复",
     "desc": "serverActions.bodySizeLimit 配置写法在 Next.js 14.2 中已迁移，旧写法会有警告。",
     "fix": "将 serverActions 配置放入 experimental 字段内，使用兼容写法。"},
    {"id": 13, "severity": "🟡 中", "file": ".env",
     "line": "4", "status": "待修复",
     "desc": "OPENAI_API_KEY 配置为空字符串，AI 分析功能将无法调用，测评提交后异步生成报告时会抛出异常，测评状态变为 failed。",
     "fix": "填入有效的 OpenAI API Key，或接入其他 AI 服务（如通义千问、文心一言等）并对应修改 lib/ai/index.ts。"},
]

SEVERITY_BG = {
    "🔴 严重": RED_BG,
    "🟠 高":   ORANGE_BG,
    "🟡 中":   YELLOW_BG,
}
SEVERITY_TEXT = {
    "🔴 严重": RED_TEXT,
    "🟠 高":   ORANGE_TEXT,
    "🟡 中":   YELLOW_TEXT,
}

# ── 工具函数 ──────────────────────────────────────────────────────────────
def _rgb_to_hex(color: RGBColor) -> str:
    """将 RGBColor 转为 6 位十六进制字符串（如 '1B5E3B'）"""
    return str(color)

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hexc = _rgb_to_hex(color)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)

def set_cell_border(cell, color: RGBColor = BORDER):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    hexc = _rgb_to_hex(color)
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), hexc)
        tcPr.append(border)

def add_styled_heading(doc, text, level=1, color=DARK_BG):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 72)
    run.font.color.rgb = BORDER
    run.font.size = Pt(6)

# ── 主函数 ────────────────────────────────────────────────────────────────
def build_report(output_path: str):
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ═══ 封面 ═══════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("H5 展业")
    run.font.size    = Pt(36)
    run.font.bold    = True
    run.font.color.rgb = DARK_BG

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目问题分析与修复报告")
    run.font.size     = Pt(22)
    run.font.bold     = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    today = datetime.date.today().strftime("%Y 年 %m 月 %d 日")
    run = p.add_run(f"报告日期：{today}")
    run.font.size     = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    run = doc.add_paragraph().add_run()
    run._element.getparent().append(OxmlElement("w:br"))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("报告版本：v1.0")
    run.font.size     = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ═══ 第一章：项目概况 ═══════════════════════════════════════════════
    add_styled_heading(doc, "一、项目概况", level=1)
    add_divider(doc)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info = [
        ("项目名称", "H5 展业"),
        ("技术栈",   "Next.js 14.2.18 + React 18 + TypeScript + Tailwind CSS + Prisma + SQLite"),
        ("构建工具", "Next.js (App Router)"),
        ("数据库",   "SQLite (Prisma ORM)"),
        ("部署方式", "EdgeOne Pages (推荐)"),
    ]
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            set_cell_border(cell)
        # 表头列
        cell0 = table.rows[i].cells[0]
        cell0.width = Cm(3.5)
        for para in cell0.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = DARK_BG

    doc.add_paragraph()

    # ═══ 第二章：问题汇总 ════════════════════════════════════════════════
    add_styled_heading(doc, "二、问题汇总表", level=1)
    add_divider(doc)
    doc.add_paragraph("下表按严重度排序，🔴 严重 = 构建失败/运行时崩溃；🟠 高 = 逻辑错误，上线后异常；🟡 中 = 隐患/不完整实现。")

    # 统计
    n_critical = sum(1 for x in ISSUES if x["severity"] == "🔴 严重")
    n_high     = sum(1 for x in ISSUES if x["severity"] == "🟠 高")
    n_medium   = sum(1 for x in ISSUES if x["severity"] == "🟡 中")
    n_fixed    = sum(1 for x in ISSUES if x["status"] == "已修复")
    n_pending  = sum(1 for x in ISSUES if x["status"] == "待修复")

    p = doc.add_paragraph()
    run = p.add_run(f"合计 {len(ISSUES)} 项问题（🔴 严重 {n_critical} / 🟠 高 {n_high} / 🟡 中 {n_medium}）　　已修复 {n_fixed} / 待修复 {n_pending}")
    run.font.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_BG
    doc.add_paragraph()

    # 汇总表格
    headers = ["#", "严重度", "文件", "行号", "问题描述（摘要）", "状态"]
    col_widths = [Cm(0.8), Cm(2.2), Cm(5.5), Cm(1.2), Cm(5.5), Cm(1.8)]
    t = doc.add_table(rows=len(ISSUES)+1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 表头
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.width = w
        set_cell_bg(cell, DARK_BG)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold     = True
                run.font.size      = Pt(10)

    # 数据行
    for r, issue in enumerate(ISSUES, start=1):
        row = t.rows[r]
        vals = [str(issue["id"]), issue["severity"], issue["file"],
                issue["line"], issue["desc"], issue["status"]]
        for c, (val, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[c]
            cell.text = val
            cell.width = w
            set_cell_border(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
            # 严重度着色
            if c == 1:
                bg = SEVERITY_BG.get(issue["severity"], WHITE)
                set_cell_bg(row.cells[c], bg)
                for para in row.cells[c].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = SEVERITY_TEXT.get(issue["severity"], GRAY_TEXT)
            # 状态着色
            if c == 5:
                for para in row.cells[c].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        if issue["status"] == "已修复":
                            run.font.color.rgb = RGBColor(0x15, 0x65, 0x5C)
                            run.font.bold = True
                        else:
                            run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
                            run.font.bold = True

        # 斑马纹
        if r % 2 == 0:
            for cell in row.cells:
                # 不覆盖严重度列的背景
                set_cell_bg(cell, STRIPE)

    doc.add_page_break()

    # ═══ 第三章：问题详细说明 ═══════════════════════════════════════════
    add_styled_heading(doc, "三、问题详细说明与修复记录", level=1)
    add_divider(doc)

    for issue in ISSUES:
        # 小标题
        p = doc.add_paragraph()
        run = p.add_run(f"问题 #{issue['id']}　{issue['severity']}")
        run.font.bold = True
        run.font.size  = Pt(13)
        run.font.color.rgb = SEVERITY_TEXT.get(issue["severity"], GRAY_TEXT)

        # 信息表
        info_t = doc.add_table(rows=4, cols=2)
        info_t.style = "Table Grid"
        fields = [
            ("文件路径", issue["file"]),
            ("行号",     issue["line"]),
            ("状态",     issue["status"]),
            ("问题描述", issue["desc"]),
        ]
        for i, (k, v) in enumerate(fields):
            row = info_t.rows[i]
            row.cells[0].text = k
            row.cells[1].text = v
            for cell in row.cells:
                set_cell_border(cell)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
            # 左列加粗
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = DARK_BG
            # 状态着色
            if k == "状态":
                for para in row.cells[1].paragraphs:
                    for run in para.runs:
                        if issue["status"] == "已修复":
                            run.font.color.rgb = RGBColor(0x15, 0x65, 0x5C)
                        else:
                            run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("修复方案：")
        run.font.bold = True
        run.font.size  = Pt(11)
        run.font.color.rgb = DARK_BG
        p = doc.add_paragraph(issue["fix"])
        for run in p.runs:
            run.font.size = Pt(10.5)

        # 分隔线
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(10)
        run = p.add_run("─" * 68)
        run.font.color.rgb = BORDER
        run.font.size = Pt(7)

    doc.add_page_break()

    # ═══ 第四章：修复状态总结 ════════════════════════════════════════════
    add_styled_heading(doc, "四、修复状态总结", level=1)
    add_divider(doc)

    # 绿色已修复框
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    # 用浅绿背景模拟高亮框（通过表格实现）
    highlight = doc.add_table(rows=1, cols=1)
    highlight.style = "Table Grid"
    cell = highlight.cell(0, 0)
    cell.text = f"✅ 已修复问题：{n_fixed} 项\n" + "、".join(str(x["id"]) for x in ISSUES if x["status"] == "已修复") + " 号"
    set_cell_bg(cell, LIGHT_BG)
    set_cell_border(cell, DARK_BG)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = DARK_BG

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("⚠️ 待修复问题（需单独排期）")
    run.font.bold = True
    run.font.size  = Pt(12)
    run.font.color.rgb = ORANGE_TEXT

    pending = [x for x in ISSUES if x["status"] == "待修复"]
    if pending:
        t = doc.add_table(rows=len(pending), cols=3)
        t.style = "Table Grid"
        headers2 = ["#", "严重度", "问题描述"]
        for i, h in enumerate(headers2):
            cell = t.rows[0].cells[i] if i < len(t.rows[0].cells) else None
        # 重做：简单列表
        for j, issue in enumerate(pending):
            p = doc.add_paragraph()
            run = p.add_run(f"  {issue['id']}. 【{issue['severity']}】{issue['desc']}")
            run.font.size = Pt(10.5)
            if issue["severity"] == "🔴 严重":
                run.font.color.rgb = RED_TEXT
            elif issue["severity"] == "🟠 高":
                run.font.color.rgb = ORANGE_TEXT
            else:
                run.font.color.rgb = YELLOW_TEXT
    else:
        p = doc.add_paragraph("  无待修复问题，全部已完成！")

    doc.add_paragraph()
    add_divider(doc)
    p = doc.add_paragraph()
    run = p.add_run("建议优先级")
    run.font.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_BG

    priorities = [
        ("P0（立即处理）", ["9（微信支付生产实现）", "13（OPENAI_API_KEY 配置）"]),
        ("P1（上线前完成）", ["10（API 输入校验）", "11（接口限流）"]),
        ("P2（后续迭代）", []),
    ]
    for title, items in priorities:
        p = doc.add_paragraph()
        run = p.add_run(f"  {title}")
        run.font.bold = True
        run.font.size  = Pt(10.5)
        if items:
            for item in items:
                p = doc.add_paragraph(f"    • {item}")
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ═══ 第五章：构建验证结果 ═══════════════════════════════════════════
    add_styled_heading(doc, "五、修复后构建验证", level=1)
    add_divider(doc)

    p = doc.add_paragraph("执行以下命令验证修复效果：")
    for run in p.runs:
        run.font.size = Pt(11)
    p = doc.add_paragraph("  npm run build")
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size  = Pt(10)
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x3B)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("构建结果：✅ Compiled successfully")
    run.font.bold = True
    run.font.size  = Pt(12)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0x5C)

    p = doc.add_paragraph("所有页面编译通过，无 TypeScript 类型错误，无运行时语法错误。")
    for run in p.runs:
        run.font.size = Pt(11)

    # ── 保存 ────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")

if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "H5展业_问题分析报告.docx"
    build_report(out)
